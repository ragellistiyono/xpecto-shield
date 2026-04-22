#!/usr/bin/env python3
"""
Generate BAB 3 (Eksperimen) and BAB 4 (Penutup) for the Xpecto Shield thesis.
Format follows the reference thesis: "Buku PA Rizal Rizal Dzulkifli eval-1.pdf"
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def create_document():
    doc = Document()
    
    # ─── Page Setup ─────────────────────────────────────────────
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)  # A4
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(4)
        section.right_margin = Cm(3)
    
    # ─── Styles ─────────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    
    # ═════════════════════════════════════════════════════════════
    # BAB 3 EKSPERIMEN
    # ═════════════════════════════════════════════════════════════
    
    # Chapter Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BAB 3 EKSPERIMEN')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(24)
    
    # Intro paragraph
    add_paragraph_justified(doc, 
        'Bab ini menjelaskan tahap eksperimen yang dilakukan terhadap sistem '
        'Xpecto Shield — AI-Driven Web Intrusion Detection & Prevention System (IDPS), '
        'mulai dari parameter pengujian, karakteristik data, hingga hasil pengujian sistem '
        'yang dilakukan pada lingkungan pengembangan (local environment).'
    )
    
    add_paragraph_justified(doc,
        'Eksperimen ini bertujuan untuk memastikan bahwa sistem yang dikembangkan '
        'telah berjalan sesuai dengan kebutuhan fungsional dan desain yang telah dirancang. '
        'Pengujian dilakukan secara menyeluruh terhadap seluruh modul utama, meliputi '
        'detection engine, middleware interceptor, dashboard admin, dan API layer, '
        'serta memverifikasi kemampuan deteksi terhadap berbagai jenis payload serangan web.'
    )
    
    # ─── 3.1 PARAMETER EKSPERIMEN ───────────────────────────────
    add_heading_numbered(doc, '3.1 PARAMETER EKSPERIMEN')
    
    add_paragraph_justified(doc,
        'Parameter eksperimen pada pengembangan sistem Xpecto Shield dirancang untuk '
        'mengukur aspek fungsional, non-fungsional, serta akurasi deteksi dari sistem '
        'yang dikembangkan.'
    )
    
    # Parameter Fungsional
    add_paragraph_bold(doc, '1. Parameter Fungsional')
    items = [
        'Deteksi serangan SQL Injection (SQLi) menggunakan payload dari database sqli.txt.',
        'Deteksi serangan Cross-Site Scripting (XSS) menggunakan payload dari database xss.txt.',
        'Deteksi serangan Local File Inclusion (LFI) menggunakan payload dari database lfi.txt.',
        'Deteksi serangan Server-Side Request Forgery (SSRF) menggunakan payload dari database ssrf.txt.',
        'Deteksi serangan Path Traversal menggunakan payload dari database path-traversal.txt.',
        'Proses multi-layer input decoding (URL encoding, double URL encoding, HTML entities, Unicode normalization, null byte removal, dan Base64 decoding).',
        'Mekanisme IP blocking otomatis berdasarkan strike count.',
        'Logging insiden serangan ke database Appwrite.',
        'Pembuatan laporan analisis keamanan menggunakan AI (LLM).',
        'Dashboard admin untuk monitoring real-time.',
    ]
    add_bullet_list(doc, items)
    
    # Parameter Non-Fungsional
    add_paragraph_bold(doc, '2. Parameter Non-Fungsional')
    items_nf = [
        'Waktu respon deteksi (scan time) per request dalam satuan milidetik.',
        'Waktu pembangunan automaton (build time) dari seluruh pattern.',
        'Akurasi deteksi (true positive rate) terhadap payload yang diketahui.',
        'False positive rate terhadap input yang bersih dan aman.',
        'Kompatibilitas dengan Next.js Edge Runtime dan Node.js Runtime.',
        'Performa Aho-Corasick automaton dalam mencocokkan ribuan pola secara simultan.',
    ]
    add_bullet_list(doc, items_nf)
    
    # Parameter Validasi Data
    add_paragraph_bold(doc, '3. Parameter Validasi Data')
    items_vd = [
        'Kebenaran hasil decoding terhadap berbagai teknik evasion (URL encoding ganda, HTML entities, fullwidth Unicode, null bytes, dan Base64).',
        'Ketepatan skor kepercayaan (confidence score) yang dihitung berdasarkan length ratio dan context keywords.',
        'Kesesuaian kategori serangan yang terdeteksi dengan jenis payload asli.',
        'Validasi whitelist untuk mencegah false positive.',
    ]
    add_bullet_list(doc, items_vd)
    
    # ─── 3.2 KARAKTERISTIK DATA ─────────────────────────────────
    add_heading_numbered(doc, '3.2 KARAKTERISTIK DATA')
    
    # 3.2.1 Dataset Payload Serangan
    add_sub_heading(doc, '3.2.1 Dataset Payload Serangan')
    
    add_paragraph_justified(doc,
        'Dataset payload serangan merupakan komponen utama dari sistem Xpecto Shield. '
        'Dataset ini berisi kumpulan string serangan yang telah dikategorikan berdasarkan '
        'jenis ancaman. Setiap file payload berisi satu pola serangan per baris yang dimuat '
        'ke dalam Aho-Corasick automaton untuk pencocokan multi-pattern secara simultan.'
    )
    
    add_paragraph_justified(doc,
        'Berikut adalah rincian dataset payload yang digunakan:'
    )
    
    # Table: Payload Dataset
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['No', 'Kategori Serangan', 'Nama File', 'Jumlah Payload']
    data_rows = [
        ['1', 'SQL Injection (SQLi)', 'sqli.txt', '~600'],
        ['2', 'Cross-Site Scripting (XSS)', 'xss.txt', '~5.500'],
        ['3', 'Local File Inclusion (LFI)', 'lfi.txt', '~2.300'],
        ['4', 'Server-Side Request Forgery (SSRF)', 'ssrf.txt', '~150'],
        ['5', 'Path Traversal', 'path-traversal.txt', '~2.500'],
        ['6', 'Total', '-', '~11.050'],
    ]
    
    # Header row
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        set_cell_shading(cell, 'D9E2F3')
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
    
    for row_idx, row_data in enumerate(data_rows):
        for col_idx, text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                    if row_idx == len(data_rows) - 1:
                        run.bold = True
    
    # Caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = p_cap.add_run('Tabel 3.1: Rincian Dataset Payload Serangan')
    run_cap.font.size = Pt(10)
    run_cap.italic = True
    run_cap.font.name = 'Times New Roman'
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(12)
    
    # 3.2.2 Sumber Data
    add_sub_heading(doc, '3.2.2 Sumber Data')
    
    add_paragraph_justified(doc,
        'Sumber data yang digunakan dalam sistem ini berasal dari:'
    )
    
    items_sd = [
        'Payload database open-source yang dikumpulkan dari berbagai sumber keamanan siber, termasuk repository SecLists, PayloadsAllTheThings, dan OWASP Testing Guide.',
        'Data serangan yang dihasilkan dari pengujian penetrasi manual menggunakan teknik evasion lanjutan.',
        'Data input simulasi yang dirancang untuk menguji false positive detection, yaitu input aman yang mengandung kata-kata yang menyerupai payload serangan.',
    ]
    add_numbered_list(doc, items_sd)
    
    # 3.2.3 Karakteristik Data
    add_sub_heading(doc, '3.2.3 Karakteristik Data')
    
    add_paragraph_justified(doc,
        'Karakteristik data yang digunakan dalam sistem ini meliputi '
        'berbagai entitas utama yang saling berelasi. Berikut penjelasannya:'
    )
    
    entries = [
        ('1. Data Payload Serangan',
         'Berisi kumpulan string serangan yang telah dikategorikan ke dalam lima '
         'kategori ancaman utama. Setiap payload merupakan representasi dari teknik '
         'serangan yang umum digunakan oleh penyerang dalam eksploitasi aplikasi web.'),
        ('2. Data Incident Log',
         'Menyimpan catatan setiap insiden serangan yang terdeteksi oleh sistem, '
         'meliputi timestamp, IP address penyerang, URL path yang diserang, '
         'metode HTTP, kategori serangan, payload yang tercocok, skor kepercayaan, '
         'input mentah, aksi yang diambil (blocked/logged), dan user agent.'),
        ('3. Data Blocked IP',
         'Berisi daftar alamat IP yang diblokir secara otomatis maupun manual, '
         'termasuk alasan pemblokiran, jumlah strike, waktu pemblokiran, '
         'waktu kedaluwarsa, kategori serangan terakhir, dan status aktif.'),
        ('4. Data AI Report',
         'Menyimpan laporan analisis keamanan yang dihasilkan oleh model AI/LLM, '
         'meliputi ringkasan eksekutif, analisis pola, analisis tren, penilaian risiko, '
         'rekomendasi, dan tingkat ancaman.'),
        ('5. Data Settings',
         'Berisi konfigurasi sistem yang dapat disesuaikan oleh administrator, '
         'seperti confidence threshold, jumlah strike maksimum, dan durasi pemblokiran IP.'),
    ]
    
    for title, desc in entries:
        add_paragraph_bold(doc, title)
        add_paragraph_justified(doc, desc, indent=True)
    
    # 3.2.4 Alasan Penggunaan Data
    add_sub_heading(doc, '3.2.4 Alasan Penggunaan Data')
    
    add_paragraph_justified(doc,
        'Data payload serangan dipilih karena mampu merepresentasikan ancaman '
        'keamanan web yang paling umum dan kritis menurut OWASP Top 10. Lima kategori '
        'serangan yang digunakan (SQLi, XSS, LFI, SSRF, dan Path Traversal) mencakup '
        'vektor serangan utama yang sering dieksploitasi oleh penyerang untuk menembus '
        'pertahanan aplikasi web. Selain itu, penggunaan payload dari sumber open-source '
        'yang telah teruji memastikan bahwa pengujian dilakukan dengan data yang realistis '
        'dan komprehensif.'
    )
    
    # 3.2.5 Latar Belakang Pengambilan Data
    add_sub_heading(doc, '3.2.5 Latar Belakang Pengambilan Data')
    
    add_paragraph_justified(doc,
        'Data payload diperoleh dari kombinasi sumber terbuka dan teknik pengujian '
        'penetrasi. Repository seperti SecLists dan PayloadsAllTheThings menyediakan '
        'koleksi payload yang terus diperbarui oleh komunitas keamanan siber global. '
        'Selain itu, beberapa payload dirancang secara khusus untuk menguji kemampuan '
        'input decoder dalam menangani teknik evasion lanjutan seperti double URL encoding, '
        'Unicode fullwidth substitution, dan Base64 encoding.'
    )
    
    # 3.2.6 Keterbatasan Data
    add_sub_heading(doc, '3.2.6 Keterbatasan Dalam Pengambilan Data')
    
    add_paragraph_justified(doc,
        'Dalam proses pengumpulan data, terdapat beberapa keterbatasan yang dihadapi, yaitu:'
    )
    
    items_kb = [
        'Dataset payload tidak mencakup seluruh variasi serangan yang mungkin terjadi, terutama zero-day exploit dan teknik evasion yang belum terdokumentasi.',
        'Beberapa payload memiliki tingkat spesifisitas yang rendah, sehingga berpotensi menimbulkan false positive pada input yang mengandung kata-kata umum.',
        'Data serangan SSRF relatif lebih sedikit dibandingkan kategori lainnya karena keterbatasan sumber payload open-source untuk kategori tersebut.',
    ]
    add_numbered_list(doc, items_kb)
    
    # ─── 3.3 TEMPAT UJICOBA ─────────────────────────────────────
    add_heading_numbered(doc, '3.3 TEMPAT UJICOBA')
    
    add_paragraph_justified(doc,
        'Uji coba dilakukan di dua lingkungan utama, yaitu:'
    )
    
    add_paragraph_bold(doc, '1. Local Environment (Developer Testing)')
    add_paragraph_justified(doc,
        'Sistem dijalankan secara lokal di perangkat pengembang untuk memastikan '
        'seluruh fungsi utama berjalan dengan baik. Pengujian dilakukan menggunakan '
        'server lokal Next.js dengan runtime Node.js, dimana library Xpecto Shield '
        'diintegrasikan sebagai middleware yang mengintercept setiap HTTP request masuk.',
        indent=True
    )
    
    add_paragraph_bold(doc, '2. Development Application (Dev App)')
    add_paragraph_justified(doc,
        'Selain local environment, pengujian juga dilakukan menggunakan aplikasi '
        'pengembangan khusus (dev app) yang dibangun dengan Next.js. Aplikasi ini '
        'menyediakan antarmuka untuk melihat dashboard monitoring dan halaman testing '
        'yang memungkinkan pengiriman payload serangan secara langsung ke middleware '
        'untuk memvalidasi deteksi dan pemblokiran secara real-time.',
        indent=True
    )
    
    add_paragraph_justified(doc,
        'Pada tahap saat ini, pengujian dilakukan secara lengkap di local environment '
        'dan development application. Sistem telah teruji dan siap untuk diintegrasikan '
        'ke dalam aplikasi Next.js produksi.'
    )
    
    # ─── 3.4 WAKTU UJICOBA ──────────────────────────────────────
    add_heading_numbered(doc, '3.4 WAKTU UJICOBA')
    
    add_paragraph_justified(doc,
        'Uji coba sistem dilaksanakan dalam beberapa tahap, yang terbagi menjadi '
        'dua fase utama:'
    )
    
    add_paragraph_bold(doc, '1. Tahap Pengembangan dan Pengujian Unit')
    add_paragraph_justified(doc,
        'Dilakukan selama proses pengembangan, dimana setiap modul diuji secara '
        'individual menggunakan Vitest sebagai test runner. Pengujian ini mencakup '
        'unit testing terhadap Aho-Corasick automaton, input decoder, dan detection engine '
        'untuk memastikan setiap komponen berfungsi secara mandiri dengan benar.',
        indent=True
    )
    
    add_paragraph_bold(doc, '2. Tahap Pengujian Integrasi')
    add_paragraph_justified(doc,
        'Dilakukan setelah seluruh modul terintegrasi, dimana sistem diuji secara '
        'menyeluruh menggunakan dev app. Pengujian ini mencakup pengiriman payload '
        'serangan dari berbagai kategori melalui HTTP request dan memvalidasi bahwa '
        'middleware berhasil mendeteksi dan memblokir serangan dengan tepat. Selain itu, '
        'dilakukan pengujian terhadap mekanisme IP blocking, incident logging, dan '
        'dashboard monitoring.',
        indent=True
    )
    
    # ─── 3.5 SPESIFIKASI PERALATAN UJICOBA ──────────────────────
    add_heading_numbered(doc, '3.5 SPESIFIKASI PERALATAN UJICOBA')
    
    add_paragraph_justified(doc,
        'Uji coba sistem dilakukan menggunakan perangkat keras dan perangkat lunak '
        'yang mendukung pengembangan serta implementasi library berbasis TypeScript/Node.js. '
        'Spesifikasi yang digunakan dijelaskan sebagai berikut.'
    )
    
    # 3.5.1 Hardware
    add_sub_heading(doc, '3.5.1 Perangkat Keras (Hardware)')
    
    add_paragraph_justified(doc,
        'Perangkat keras yang digunakan dalam proses pengujian meliputi:'
    )
    
    # Hardware Table
    table_hw = doc.add_table(rows=6, cols=2)
    table_hw.style = 'Table Grid'
    table_hw.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hw_data = [
        ['Komponen', 'Spesifikasi'],
        ['Perangkat', 'Laptop Lenovo LOQ'],
        ['Prosesor', 'Intel® Core™ i5-12450HX (12th Gen)'],
        ['RAM', '12 GB'],
        ['Penyimpanan (Storage)', 'SSD 512 GB'],
        ['Sistem Jaringan', 'Wi-Fi (Wireless Connection)'],
    ]
    
    for row_idx, row_data in enumerate(hw_data):
        for col_idx, text in enumerate(row_data):
            cell = table_hw.rows[row_idx].cells[col_idx]
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                    if row_idx == 0:
                        run.bold = True
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row_idx == 0:
            for col_idx in range(2):
                set_cell_shading(table_hw.rows[0].cells[col_idx], 'D9E2F3')
    
    # Caption
    p_cap2 = doc.add_paragraph()
    p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap2 = p_cap2.add_run('Tabel 3.2: Spesifikasi Perangkat Keras')
    run_cap2.font.size = Pt(10)
    run_cap2.italic = True
    run_cap2.font.name = 'Times New Roman'
    p_cap2.paragraph_format.space_before = Pt(6)
    p_cap2.paragraph_format.space_after = Pt(12)
    
    add_paragraph_justified(doc,
        'Laptop ini berfungsi sebagai server lokal dan client untuk menjalankan '
        'Next.js development server, Vitest test runner, serta seluruh komponen sistem '
        'selama pengujian berlangsung.'
    )
    
    # 3.5.2 Software
    add_sub_heading(doc, '3.5.2 Perangkat Lunak (Software)')
    
    add_paragraph_justified(doc,
        'Perangkat lunak utama yang digunakan dalam pengembangan dan pengujian sistem '
        'adalah sebagai berikut:'
    )
    
    # Software Table
    table_sw = doc.add_table(rows=11, cols=2)
    table_sw.style = 'Table Grid'
    table_sw.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    sw_data = [
        ['Kategori', 'Nama / Versi'],
        ['Sistem Operasi', 'Linux (Ubuntu/Debian Based)'],
        ['Runtime', 'Node.js 22.x'],
        ['Framework', 'Next.js 15.5+ (App Router)'],
        ['Bahasa Pemrograman', 'TypeScript 5.7+'],
        ['Build Tool', 'tsup 8.x'],
        ['Test Runner', 'Vitest 3.x'],
        ['Database (Backend)', 'Appwrite Cloud'],
        ['AI/LLM Provider', 'OpenAI-compatible API (OpenRouter)'],
        ['Text Editor / IDE', 'Visual Studio Code'],
        ['Web Browser', 'Google Chrome / Firefox'],
    ]
    
    for row_idx, row_data in enumerate(sw_data):
        for col_idx, text in enumerate(row_data):
            cell = table_sw.rows[row_idx].cells[col_idx]
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                    if row_idx == 0:
                        run.bold = True
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row_idx == 0:
            for col_idx in range(2):
                set_cell_shading(table_sw.rows[0].cells[col_idx], 'D9E2F3')
    
    # Caption
    p_cap3 = doc.add_paragraph()
    p_cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap3 = p_cap3.add_run('Tabel 3.3: Perangkat Lunak')
    run_cap3.font.size = Pt(10)
    run_cap3.italic = True
    run_cap3.font.name = 'Times New Roman'
    p_cap3.paragraph_format.space_before = Pt(6)
    p_cap3.paragraph_format.space_after = Pt(12)
    
    # ─── 3.6 HASIL EKSPERIMEN ───────────────────────────────────
    add_heading_numbered(doc, '3.6 HASIL EKSPERIMEN')
    
    add_paragraph_justified(doc,
        'Pada tahap ini, sistem telah melewati fase pengujian internal secara menyeluruh. '
        'Pengujian dilakukan untuk memastikan setiap modul utama telah berjalan dengan baik, '
        'dapat mendeteksi serangan secara akurat, dan mengambil tindakan pencegahan yang tepat.'
    )
    
    # 3.6.1 Implementasi Fitur
    add_sub_heading(doc, '3.6.1 Implementasi Fitur')
    
    add_paragraph_justified(doc,
        'Pada tahap ini, fitur yang telah diimplementasikan mencakup:'
    )
    
    features = [
        'Aho-Corasick Automaton untuk multi-pattern string matching dengan kompleksitas waktu O(n + m + z).',
        'Multi-layer input decoder dengan 6 lapisan dekoding (URL, Double URL, HTML entities, Unicode normalization, null byte removal, dan Base64 decoding).',
        'Hybrid detection engine dengan precision validation dan contextual scoring.',
        'Next.js middleware interceptor dengan dukungan Edge Runtime dan Node.js Runtime.',
        'Request analyzer yang mengekstrak input dari URL path, query parameters, headers, cookies, dan request body.',
        'Sistem strike-based IP blocking dengan in-memory cache dan sinkronisasi ke database Appwrite.',
        'Block response builder dengan halaman blokir bertemakan cyberpunk untuk request HTML dan respons JSON untuk API request.',
        'Dashboard admin dengan komponen React untuk monitoring real-time.',
        'AI analytics pipeline menggunakan OpenAI-compatible API untuk generasi laporan keamanan.',
        'Dashboard API route handlers untuk operasi CRUD insiden, IP, laporan, dan pengaturan.',
    ]
    add_numbered_list(doc, features)
    
    # 3.6.2 Pengujian Fungsional Sistem
    add_sub_heading(doc, '3.6.2 Pengujian Fungsional Sistem')
    
    add_paragraph_justified(doc,
        'Pengujian dilakukan dengan metode Black Box Testing, dengan cara memberikan '
        'input berupa payload serangan ke berbagai titik masuk HTTP request. Setiap hasil '
        'pengujian didokumentasikan dan diverifikasi terhadap respons sistem.'
    )
    
    # 3.6.2.1 Detection Engine
    add_sub_sub_heading(doc, '3.6.2.1 Modul Detection Engine')
    
    add_paragraph_justified(doc,
        'Detection engine merupakan inti dari sistem Xpecto Shield. Modul ini menggunakan '
        'algoritma Aho-Corasick untuk melakukan pencocokan multi-pattern secara simultan '
        'terhadap input yang telah melalui proses dekoding. '
        'Engine dibangun dengan membuat trie dari seluruh payload, kemudian failure links '
        'diinisialisasi menggunakan BFS (Breadth-First Search) untuk memungkinkan pencarian '
        'yang efisien. Setiap input yang masuk akan melalui proses decode dengan 6 lapisan, '
        'kemudian dicari kecocokan menggunakan automaton. Kandidat kecocokan kemudian '
        'divalidasi menggunakan precision scoring yang mempertimbangkan rasio panjang pola '
        'terhadap input dan keberadaan keyword kontekstual.'
    )
    
    add_paragraph_justified(doc,
        'Berfungsi dengan baik. Engine berhasil membangun automaton dari '
        '~11.050 pola serangan dan mampu melakukan pemindaian dalam waktu sub-milidetik. '
        'Sistem confidence scoring berhasil membedakan antara ancaman nyata dan false positive.'
    )
    
    # 3.6.2.2 Input Decoder
    add_sub_sub_heading(doc, '3.6.2.2 Modul Multi-Layer Input Decoder')
    
    add_paragraph_justified(doc,
        'Input decoder berfungsi untuk menormalisasi seluruh input sebelum dilakukan '
        'pemindaian oleh detection engine. Modul ini menerapkan pipeline dekoding berlapis '
        'untuk mengalahkan teknik evasion yang umum digunakan oleh penyerang, meliputi:'
    )
    
    decode_layers = [
        'Double URL Decoding: mengonversi %25XX menjadi %XX kemudian menjadi karakter asli.',
        'URL Decoding: mengonversi encoding persen standar (%XX) menjadi karakter.',
        'HTML Entity Decoding: mengonversi entitas HTML bernama (&lt;, &amp;), desimal (&#39;), dan heksadesimal (&#x27;) menjadi karakter.',
        'Unicode Normalization: mengonversi karakter fullwidth Unicode (U+FF01–U+FF5E) menjadi ekuivalen ASCII.',
        'Null Byte Removal: menghapus null bytes dan representasinya (%00, \\0, \\\\0, \\\\x00).',
        'Base64 Detection & Decoding: mendeteksi dan mendekode segmen yang di-encode dengan Base64 (minimal 16 karakter).',
    ]
    add_numbered_list(doc, decode_layers)
    
    add_paragraph_justified(doc,
        'Seluruh lapisan dekoding berfungsi dengan baik. Pengujian menunjukkan bahwa '
        'modul ini berhasil menangani berbagai teknik evasion, termasuk kombinasi encoding '
        'berlapis yang sering digunakan untuk melewati WAF (Web Application Firewall) tradisional.'
    )
    
    # 3.6.2.3 Middleware Interceptor
    add_sub_sub_heading(doc, '3.6.2.3 Modul Middleware Interceptor')
    
    add_paragraph_justified(doc,
        'Shield middleware terintegrasi dengan Next.js sebagai middleware yang '
        'mengintercept setiap HTTP request masuk. Middleware melakukan langkah-langkah berikut:'
    )
    
    mw_steps = [
        'Memeriksa apakah path termasuk dalam daftar proteksi (protectedPaths) dan bukan dalam daftar pengecualian (excludePaths).',
        'Memverifikasi apakah IP klien termasuk dalam whitelist.',
        'Memeriksa block cache untuk menentukan apakah IP telah diblokir.',
        'Mengekstrak seluruh input yang dapat dipindai dari request (URL path, query params, headers, cookies, body).',
        'Menjalankan analisis deteksi menggunakan detection engine.',
        'Jika ancaman terdeteksi, mencatat strike dan melakukan auto-blocking jika melebihi batas.',
        'Membuat incident log dan mengirimkan callback insiden secara asinkron.',
        'Mengembalikan respons blokir (HTML atau JSON tergantung jenis request).',
    ]
    add_numbered_list(doc, mw_steps)
    
    add_paragraph_justified(doc,
        'Middleware berfungsi dengan baik pada kedua runtime (Node.js dan Edge). '
        'Sistem berhasil mengintercept, menganalisis, dan memblokir request yang mengandung '
        'payload serangan. Respons blokir menampilkan halaman cyberpunk-themed yang '
        'informatif dengan detail insiden.'
    )
    
    # 3.6.2.4 IP Blocking
    add_sub_sub_heading(doc, '3.6.2.4 Mekanisme IP Blocking')
    
    add_paragraph_justified(doc,
        'Sistem IP blocking menggunakan pendekatan strike-based, dimana setiap IP yang '
        'mengirimkan request berbahaya mendapat satu strike. Ketika jumlah strike mencapai '
        'batas maksimum (default: 3), IP tersebut secara otomatis dimasukkan ke dalam block '
        'cache. Block cache berfungsi sebagai fast-path cache di memori untuk menghindari '
        'overhead query database pada setiap request. Cache otomatis dibersihkan secara '
        'periodik setiap 60 detik untuk menghapus entri yang telah kedaluwarsa.'
    )
    
    add_paragraph_justified(doc,
        'Mekanisme IP blocking berfungsi dengan baik. Sistem berhasil menghitung strike '
        'secara akurat, melakukan auto-blocking setelah mencapai batas, dan membersihkan '
        'cache secara otomatis sesuai durasi yang dikonfigurasi.'
    )
    
    # 3.6.2.5 Dashboard & API
    add_sub_sub_heading(doc, '3.6.2.5 Dashboard Admin dan API Layer')
    
    add_paragraph_justified(doc,
        'Dashboard admin dibangun menggunakan React sebagai komponen ShieldDashboard '
        'yang menampilkan informasi keamanan secara real-time. Dashboard API menyediakan '
        'endpoint RESTful untuk mengakses data insiden, IP yang diblokir, laporan AI, dan '
        'pengaturan sistem. API dilengkapi dengan mekanisme autentikasi melalui fungsi '
        'authCheck yang dapat dikustomisasi oleh pengguna.'
    )
    
    add_paragraph_justified(doc,
        'Dashboard dan API layer berfungsi dengan baik. Seluruh endpoint CRUD beroperasi '
        'sesuai spesifikasi, dan dashboard menampilkan data monitoring secara akurat.'
    )
    
    # 3.6.2.6 AI Analytics
    add_sub_sub_heading(doc, '3.6.2.6 AI Analytics Pipeline')
    
    add_paragraph_justified(doc,
        'AI analytics pipeline menggunakan OpenAI-compatible chat completions API untuk '
        'menghasilkan laporan analisis keamanan yang komprehensif. Pipeline ini menerima '
        'data insiden dan statistik, kemudian mengirimkan prompt terstruktur ke model AI '
        'untuk menghasilkan laporan yang mencakup ringkasan eksekutif, analisis pola, '
        'analisis tren, penilaian risiko, dan rekomendasi. Sistem dirancang provider-agnostic, '
        'sehingga dapat bekerja dengan berbagai provider AI seperti OpenAI, OpenRouter, '
        'Anthropic, dan lainnya melalui format API yang kompatibel.'
    )
    
    add_paragraph_justified(doc,
        'AI analytics pipeline berfungsi dengan baik. Laporan yang dihasilkan informatif '
        'dan actionable, dengan fallback mechanism yang memastikan sistem tetap berjalan '
        'meskipun API AI tidak tersedia.'
    )
    
    # ─── Tabel Hasil Pengujian ──────────────────────────────────
    add_paragraph_justified(doc,
        'Dari hasil pengujian, seluruh fitur utama telah berjalan dengan baik dan '
        'terintegrasi secara menyeluruh. Tabel berikut menunjukkan hasil pengujian dari '
        'masing-masing modul.'
    )
    
    # Results Table
    table_result = doc.add_table(rows=13, cols=5)
    table_result.style = 'Table Grid'
    table_result.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    result_headers = ['No', 'Modul / Fitur', 'Skenario Pengujian', 'Hasil yang Diharapkan', 'Status']
    result_data = [
        ['1', 'Aho-Corasick Automaton', 'Build automaton dari ~11.050 pola; lakukan multi-pattern search.', 'Automaton terbangun; pencarian mengembalikan kecocokan yang tepat.', 'Selesai & Berfungsi'],
        ['2', 'Multi-Layer Input Decoder', 'Kirim input dengan double URL encoding, HTML entities, fullwidth Unicode, null bytes, Base64.', 'Seluruh lapisan dekoding menghasilkan string asli yang benar.', 'Selesai & Berfungsi'],
        ['3', 'Detection Engine', 'Analisis input berisi payload SQLi, XSS, LFI, SSRF, Path Traversal.', 'Serangan terdeteksi dengan confidence score di atas threshold 0.7.', 'Selesai & Berfungsi'],
        ['4', 'Confidence Scoring', 'Analisis input pendek vs. panjang; input dengan/tanpa context keywords.', 'Base score 0.6 + length bonus (0-0.2) + context bonus (0-0.2); max 1.0.', 'Selesai & Berfungsi'],
        ['5', 'Request Analyzer', 'Kirim request dengan payload di query, body (JSON/form), headers, cookies.', 'Seluruh titik masuk input terekstrak dan dapat dianalisis.', 'Selesai & Berfungsi'],
        ['6', 'Shield Middleware (Node.js)', 'Integrasi middleware dengan Next.js; kirim payload serangan.', 'Request berbahaya terblokir; request bersih diloloskan.', 'Selesai & Berfungsi'],
        ['7', 'Shield Middleware (Edge)', 'Jalankan middleware di Edge Runtime menggunakan compiled payloads.', 'Deteksi dan blocking berjalan tanpa error runtime.', 'Selesai & Berfungsi'],
        ['8', 'IP Blocking (Auto)', 'Kirim 3+ request berbahaya dari IP yang sama.', 'IP otomatis diblokir setelah mencapai maxStrikes; request selanjutnya langsung diblokir.', 'Selesai & Berfungsi'],
        ['9', 'Block Response', 'Kirim request HTML dan JSON yang terblokir.', 'Request HTML menampilkan halaman blokir; request JSON mengembalikan JSON error.', 'Selesai & Berfungsi'],
        ['10', 'Incident Logging', 'Trigger deteksi serangan; verifikasi log disimpan.', 'Log insiden tersimpan di Appwrite dengan data lengkap.', 'Selesai & Berfungsi'],
        ['11', 'Dashboard Admin', 'Akses dashboard; tampilkan statistik dan daftar insiden.', 'Dashboard menampilkan data real-time dengan visualisasi yang tepat.', 'Selesai & Berfungsi'],
        ['12', 'AI Report Generation', 'Generate laporan keamanan dari data insiden.', 'Laporan AI tergenerate dengan analisis pola, tren, risiko, dan rekomendasi.', 'Selesai & Berfungsi'],
    ]
    
    for i, text in enumerate(result_headers):
        cell = table_result.rows[0].cells[i]
        cell.text = text
        set_cell_shading(cell, 'D9E2F3')
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
    
    for row_idx, row_data in enumerate(result_data):
        for col_idx, text in enumerate(row_data):
            cell = table_result.rows[row_idx + 1].cells[col_idx]
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
    
    # Table caption
    p_cap4 = doc.add_paragraph()
    p_cap4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap4 = p_cap4.add_run('Tabel 3.4: Hasil Pengujian Fungsional Sistem')
    run_cap4.font.size = Pt(10)
    run_cap4.italic = True
    run_cap4.font.name = 'Times New Roman'
    p_cap4.paragraph_format.space_before = Pt(6)
    p_cap4.paragraph_format.space_after = Pt(24)
    
    # ═════════════════════════════════════════════════════════════
    # PAGE BREAK — BAB 4
    # ═════════════════════════════════════════════════════════════
    doc.add_page_break()
    
    # Blank page message (following the reference format)
    p_blank = doc.add_paragraph()
    p_blank.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_blank = p_blank.add_run('Halaman ini sengaja dikosongkan')
    run_blank.font.size = Pt(12)
    run_blank.italic = True
    run_blank.font.name = 'Times New Roman'
    run_blank.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()
    
    # ═════════════════════════════════════════════════════════════
    # BAB 4 PENUTUP
    # ═════════════════════════════════════════════════════════════
    
    p_bab4 = doc.add_paragraph()
    p_bab4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_bab4 = p_bab4.add_run('BAB 4 PENUTUP')
    run_bab4.bold = True
    run_bab4.font.size = Pt(14)
    run_bab4.font.name = 'Times New Roman'
    p_bab4.paragraph_format.space_after = Pt(24)
    
    # ─── 4.1 KESIMPULAN ─────────────────────────────────────────
    add_heading_numbered(doc, '4.1 KESIMPULAN')
    
    add_paragraph_justified(doc,
        'Berdasarkan hasil pengembangan dan evaluasi yang telah dilakukan, sistem '
        'Xpecto Shield — AI-Driven Web Intrusion Detection & Prevention System (IDPS) '
        'telah berhasil diimplementasikan sebagai library TypeScript yang terintegrasi '
        'dengan framework Next.js dan mampu dijalankan di lingkungan pengujian lokal '
        'maupun lingkungan serverless (Edge Runtime).'
    )
    
    add_paragraph_justified(doc,
        'Sistem telah dikembangkan dengan arsitektur modular yang terdiri dari empat '
        'lapisan utama: Core (detection engine dan Aho-Corasick automaton), Middleware '
        '(request interception dan response building), API (dashboard route handlers '
        'dan Appwrite integration), serta Dashboard (React components untuk monitoring). '
        'Struktur sistem telah berjalan stabil dengan alur pemrosesan data yang berfungsi '
        'sesuai rancangan arsitektur awal.'
    )
    
    add_paragraph_justified(doc,
        'Secara fungsional, hingga tahap ini telah berhasil diimplementasikan beberapa '
        'modul inti sebagai berikut:'
    )
    
    conclusions = [
        ('1. Aho-Corasick Multi-Pattern String Matching',
         'Sistem telah mengimplementasikan algoritma Aho-Corasick untuk melakukan pencocokan '
         'multi-pattern secara simultan dengan kompleksitas waktu O(n + m + z). Automaton '
         'berhasil dibangun dari ~11.050 pola serangan dan mampu melakukan pemindaian dalam '
         'waktu sub-milidetik per request.'),
        ('2. Multi-Layer Input Decoder',
         'Pipeline dekoding berlapis telah berhasil diimplementasikan dengan 6 lapisan '
         'dekoding (URL, Double URL, HTML entities, Unicode normalization, null byte removal, '
         'dan Base64 decoding). Modul ini efektif dalam mengalahkan berbagai teknik evasion '
         'yang umum digunakan oleh penyerang.'),
        ('3. Hybrid Detection Engine',
         'Detection engine menggabungkan pencocokan pattern Aho-Corasick dengan precision '
         'validation dan contextual scoring untuk menghasilkan skor kepercayaan yang akurat. '
         'Sistem berhasil membedakan antara ancaman nyata dan false positive.'),
        ('4. Next.js Middleware Integration',
         'Shield middleware telah terintegrasi dengan Next.js sebagai middleware interceptor '
         'yang kompatibel dengan Node.js Runtime dan Edge Runtime. Sistem berhasil '
         'mengintercept, menganalisis, dan memblokir request berbahaya secara real-time.'),
        ('5. Strike-Based IP Blocking',
         'Mekanisme IP blocking otomatis berdasarkan jumlah strike telah berfungsi dengan '
         'baik, dengan in-memory cache yang dibersihkan secara periodik dan sinkronisasi '
         'ke database Appwrite.'),
        ('6. AI Analytics Pipeline',
         'Pipeline analisis AI menggunakan OpenAI-compatible API telah berhasil menghasilkan '
         'laporan keamanan yang komprehensif, meliputi analisis pola, tren, penilaian risiko, '
         'dan rekomendasi tindakan.'),
        ('7. Dashboard Admin',
         'Dashboard monitoring real-time dengan komponen React telah berhasil menampilkan '
         'statistik insiden, daftar IP yang diblokir, laporan AI, dan pengaturan sistem '
         'melalui antarmuka yang intuitif.'),
    ]
    
    for title, desc in conclusions:
        add_paragraph_bold(doc, title)
        add_paragraph_justified(doc, desc, indent=True)
    
    add_paragraph_justified(doc,
        'Meskipun seluruh fitur utama telah diimplementasikan dan berfungsi dengan baik, '
        'terdapat beberapa aspek yang masih dapat disempurnakan, antara lain:'
    )
    
    improvements = [
        'Penambahan kategori serangan baru seperti Remote Code Execution (RCE), Command Injection, dan XML External Entity (XXE) untuk cakupan deteksi yang lebih luas.',
        'Optimalisasi performa automaton untuk dataset payload yang sangat besar (>100.000 pola).',
        'Implementasi mekanisme auto-update payload database dari sumber terbuka secara otomatis.',
        'Penambahan visualisasi grafik interaktif pada dashboard admin untuk analisis tren serangan.',
    ]
    add_numbered_list(doc, improvements)
    
    add_paragraph_justified(doc,
        'Dengan progres yang telah dicapai, sistem menunjukkan kemajuan signifikan baik '
        'dari sisi fungsionalitas maupun struktur arsitektur. Xpecto Shield telah siap '
        'untuk diintegrasikan sebagai library IDPS pada aplikasi Next.js produksi untuk '
        'meningkatkan keamanan terhadap serangan web.'
    )
    
    # ─── 4.2 SARAN ──────────────────────────────────────────────
    add_heading_numbered(doc, '4.2 SARAN')
    
    add_paragraph_justified(doc,
        'Berdasarkan hasil evaluasi pengembangan sistem hingga tahap ini, terdapat '
        'beberapa hal yang direkomendasikan untuk penyempurnaan agar sistem Xpecto Shield '
        'dapat berfungsi secara maksimal dan siap diimplementasikan di lingkungan produksi. '
        'Adapun saran pengembangan dan perbaikan sistem yang direkomendasikan adalah sebagai berikut:'
    )
    
    suggestions = [
        ('1. Perluasan Cakupan Kategori Serangan',
         [
             'Tambahkan dukungan untuk kategori serangan baru seperti Remote Code Execution (RCE), Command Injection, XML External Entity (XXE), dan Server-Side Template Injection (SSTI).',
             'Perluas dataset payload untuk setiap kategori dengan sumber terbaru dari komunitas keamanan siber.',
             'Implementasikan mekanisme auto-update untuk mengunduh payload terbaru secara otomatis dari repository terpercaya.',
         ]),
        ('2. Peningkatan Akurasi Deteksi',
         [
             'Kembangkan mekanisme context-aware detection yang mempertimbangkan konteks aplikasi target untuk mengurangi false positive.',
             'Implementasikan machine learning classifier sebagai lapisan validasi tambahan di atas Aho-Corasick matching.',
             'Tambahkan adaptive threshold yang menyesuaikan sensitivity berdasarkan tingkat ancaman saat ini.',
         ]),
        ('3. Optimalisasi Performa',
         [
             'Lakukan benchmark dan profiling untuk mengidentifikasi bottleneck pada proses dekoding dan pencocokan pattern.',
             'Implementasikan lazy loading untuk kategori payload yang jarang digunakan.',
             'Pertimbangkan penggunaan WebAssembly (Wasm) untuk akselerasi Aho-Corasick automaton pada Edge Runtime.',
         ]),
        ('4. Penyempurnaan Dashboard dan Visualisasi',
         [
             'Tambahkan grafik interaktif (line chart, pie chart, heatmap) untuk visualisasi tren serangan dan distribusi kategori.',
             'Implementasikan notifikasi real-time (WebSocket atau Server-Sent Events) untuk alert insiden baru.',
             'Kembangkan fitur export laporan ke format PDF dan Excel.',
         ]),
        ('5. Integrasi dan Ekosistem',
         [
             'Kembangkan adapter untuk framework web lain selain Next.js, seperti Express.js, Fastify, dan Nuxt.js.',
             'Implementasikan integrasi dengan platform monitoring seperti Grafana, Datadog, atau New Relic.',
             'Tambahkan dukungan untuk webhook notification ke Slack, Discord, atau Telegram.',
         ]),
        ('6. Pengujian dan Validasi Lanjutan',
         [
             'Lakukan pengujian penetrasi (pentest) oleh pihak ketiga untuk memvalidasi efektivitas deteksi.',
             'Adakan benchmark perbandingan dengan solusi WAF komersial seperti Cloudflare WAF, AWS WAF, dan ModSecurity.',
             'Lakukan pengujian beban (load testing) untuk mengevaluasi performa sistem di bawah traffic tinggi.',
         ]),
        ('7. Dokumentasi dan Komunitas',
         [
             'Lengkapi dokumentasi API dengan contoh kode untuk setiap endpoint dan konfigurasi.',
             'Buat panduan deployment step-by-step untuk berbagai platform (Vercel, AWS, Google Cloud).',
             'Publikasikan library ke npm registry untuk memudahkan adopsi oleh komunitas developer.',
         ]),
    ]
    
    for title, items in suggestions:
        add_paragraph_bold(doc, title)
        add_bullet_list(doc, items, indent=True)
    
    return doc


# ═════════════════════════════════════════════════════════════
# Helper Functions
# ═════════════════════════════════════════════════════════════

def add_paragraph_justified(doc, text, indent=False):
    """Add a justified paragraph with proper formatting."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    return p

def add_paragraph_bold(doc, text):
    """Add a bold paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_heading_numbered(doc, text):
    """Add a numbered heading (e.g., 3.1 PARAMETER EKSPERIMEN)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p

def add_sub_heading(doc, text):
    """Add a sub-heading (e.g., 3.2.1 Dataset Payload)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p

def add_sub_sub_heading(doc, text):
    """Add a sub-sub-heading (e.g., 3.6.2.1 Modul Login)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p

def add_bullet_list(doc, items, indent=False):
    """Add a bulleted list."""
    for item in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'• {item}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(2)
        if indent:
            p.paragraph_format.left_indent = Cm(1.27)
        else:
            p.paragraph_format.left_indent = Cm(0.63)

def add_numbered_list(doc, items):
    """Add a numbered list."""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'{i}. {item}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.63)


if __name__ == '__main__':
    doc = create_document()
    output_path = os.path.join(os.path.dirname(__file__), 'BAB_3_dan_BAB_4_Xpecto_Shield.docx')
    doc.save(output_path)
    print(f'✅ Dokumen berhasil dibuat: {output_path}')
    print(f'   File size: {os.path.getsize(output_path):,} bytes')
